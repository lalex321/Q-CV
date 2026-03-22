import os
import sys
import json
import re
import warnings
import copy
import platform
import subprocess
import time
import shutil
import datetime
def _extract_contacts_plus(raw_text: str) -> dict:
    """Best-effort extraction of email/phone/website/linkedin from raw text."""
    out = {"email":"", "phone":"", "website":"", "linkedin":""}
    if not isinstance(raw_text, str) or not raw_text.strip():
        return out
    t = raw_text

    m = re.search(r"([A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,})", t, re.I)
    if m: out["email"] = m.group(1)

    for m in re.finditer(r"(\+?\d[\d\s().-]{7,}\d)", t):
        cand = m.group(1)
        digits = re.sub(r"\D", "", cand)
        if len(digits) >= 9:
            out["phone"] = cand.strip()
            break

    m = re.search(r"(https?://(?:[a-z]{2,3}\.)?linkedin\.com/[\w\-./?=&%#]+)", t, re.I)
    if m: out["linkedin"] = m.group(1)

    m = re.search(r"(https?://(?![^\s]*linkedin\.com)[^\s)]+)", t, re.I)
    if m:
        out["website"] = m.group(1)
    else:
        m = re.search(r"\b(www\.[^\s)]+)", t, re.I)
        if m:
            out["website"] = m.group(1)

    return out

def _extract_location_line(raw_text: str) -> str:
    """Try to find a short location-like line near the top."""
    if not isinstance(raw_text, str) or not raw_text.strip():
        return ""
    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]
    for ln in lines[:25]:
        if len(ln) > 60:
            continue
        low = ln.lower()
        if any(k in low for k in ("summary","experience","education","skills","top skills","certification","linkedin","http","www","@", "project:")):
            continue
        if "," in ln:
            return ln
    return ""



# --- 1. BASIC SETTINGS ---


def _trim_strings_deep(value):
    """Recursively trim whitespace around strings before rendering/saving."""
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, list):
        return [_trim_strings_deep(v) for v in value]
    if isinstance(value, dict):
        return {k: _trim_strings_deep(v) for k, v in value.items()}
    return value

warnings.filterwarnings("ignore")
os.environ["GRPC_VERBOSITY"] = "ERROR"
os.environ["TF_CPP_MIN_LOG_LEVEL"] = "3"

from google import genai
from google.genai import types as genai_types

# ==========================================
# 🛡️ GLOBAL API RETRY HELPER (429 ERROR FIX)
# ==========================================
def _retry_generate(client, model_name, contents):
    max_retries = 3
    delay = 5
    for attempt in range(max_retries):
        try:
            return client.models.generate_content(model=model_name, contents=contents)
        except Exception as e:
            err_str = str(e)
            if "429" in err_str or "Resource exhausted" in err_str or "Quota" in err_str:
                if attempt < max_retries - 1:
                    print(f"⚠️ API 429 Limit hit. Sleeping for {delay} seconds... (Attempt {attempt+1}/{max_retries})")
                    time.sleep(delay)
                else:
                    raise e
            else:
                raise e
# ==========================================

try:
    from docxtpl import DocxTemplate
except ImportError:
    print("CRITICAL: docxtpl not installed. Please run: pip install docxtpl")

# ==========================================
# 🛑 DOCX "GHOST FOLDER" FIX
# ==========================================
def fix_docx_path_bug():
    if getattr(sys, 'frozen', False):
        try:
            base_dir = sys._MEIPASS
            parts_dir = os.path.join(base_dir, 'docx', 'parts')
            os.makedirs(parts_dir, exist_ok=True)
            templates_dir = os.path.join(base_dir, 'docx', 'templates')
            os.makedirs(templates_dir, exist_ok=True)
            header = os.path.join(templates_dir, 'default-header.xml')
            if not os.path.exists(header):
                with open(header, 'wb') as f:
                    f.write(b"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n<w:hdr xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"></w:hdr>")
            footer = os.path.join(templates_dir, 'default-footer.xml')
            if not os.path.exists(footer):
                with open(footer, 'wb') as f:
                    f.write(b"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n<w:ftr xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"></w:ftr>")
        except: pass

fix_docx_path_bug()
# ==========================================

MODEL_NAME = 'gemini-2.0-flash'
APP_VERSION = "03.50"

# 💸 GEMINI 2.0 FLASH PRICE LIST (Per 1 Million Tokens)
PRICE_1M_IN = 0.15
PRICE_1M_OUT = 0.60

def get_resource_path(relative_path):
    if getattr(sys, 'frozen', False):
        try:
            base_path = sys._MEIPASS
            if os.path.exists(os.path.join(base_path, relative_path)):
                return os.path.join(base_path, relative_path)
        except Exception: pass
        app_resources_path = os.path.join(os.path.dirname(sys.executable), '..', 'Resources')
        if os.path.exists(os.path.join(app_resources_path, relative_path)):
            return os.path.join(app_resources_path, relative_path)
    return os.path.join(os.path.abspath(os.path.dirname(__file__)), relative_path)

USER_HOME = os.path.expanduser("~")
DEFAULT_WORKSPACE = os.path.join(USER_HOME, "Documents", "Quantori_CV_Workplace")
SCRIPT_NAME = "QuantoriCV" if getattr(sys, 'frozen', False) else "CV Manager"
SETTINGS_FILE = os.path.join(USER_HOME, '.quantoricv_settings.json')
MASTER_PROMPTS_FILE = os.path.join(USER_HOME, '.master_prompts.json')

# ==========================================
# 2. PROMPTS & SCHEMAS (PROMPT LAB FOUNDATION)
# ==========================================

# 🛡️ PROTECTED SCHEMA: Never exposed to the user for editing
CV_JSON_SCHEMA = """{
  "basics": {
    "name": "String", "current_title": "String", "objective": "String",
    "contacts": { "email": "String", "phone": "String", "location": "String" },
    "links": ["String"]
  },
  "summary": { "bullet_points": ["String"] },
  "skills": { "Category": ["Skill 1"] },
  "experience": [{
      "category": "String", "company_name": "String", "role": "String",
      "dates": { "start": "String", "end": "String" },
      "location": "String", "project_description": "String",
      "highlights": ["String"], "environment": ["String"]
  }],
  "education": [{"institution": "String", "degree": "String", "year": "String", "details": "String"}],
  "certifications": ["String"],
  "languages": [{"language": "String", "proficiency": "String", "level": "String", "details": "String"}],
  "other_sections": [{"title": "String", "items": ["String"]}]
}"""

# 🧠 MUTABLE PROMPTS: Available for editing in the UI
DEFAULT_PROMPTS = {
    "prompt_master_inst": """You are a STRICT Lossless CV Extractor and Translator. Your ONLY job is to digitize this CV into JSON matching the required schema with minimal data loss.

**CRITICAL RULES FOR LOSSLESS EXTRACTION (STRICTLY ENFORCED):**
1. **US ENGLISH ONLY:** All human-readable output in the JSON must be translated into professional US English. Do not leave Russian or other non-English prose in the output.
2. **NO INVENTED FACTS / NO DATA LOSS:** Extract only facts supported by the CV, but do not lose explicit information. Preserve meaningful technical terms, methods, tools, technologies, responsibilities, achievements, project details, and bullet points.
3. **DEEP SCAN THE ENTIRE CV:** Extract from the whole document, including header, summary, skills blocks, Top Skills, experience bullets, project descriptions, certifications, languages, links, and side sections.
4. **SKILLS & ENVIRONMENT:** Extract explicit skills, tools, technologies, frameworks, platforms, databases, cloud/services, and domain systems from all relevant sections. Put them into `skills` or role `environment` as appropriate. Do not create noisy or redundant generic skill categories.
5. **DATES & CURRENT STATUS:** Extract all explicit dates from all sections, preserving the highest precision supported by the source (`April 2025`, `2018`, `Present`). Keep explicit durations when present. Never invent dates. Never assign future dates to finished past roles.
6. **CONTACTS, LINKS, LOCATIONS:** Extract all explicit phone numbers, emails, LinkedIn, GitHub, portfolio, websites, WhatsApp, and other links, plus the most granular explicit location. Do not infer location from vague context or company headquarters.
7. **WORK EXPERIENCE INTEGRITY:** Merge all employment history into the single `experience` array, even if the CV splits it into multiple employment sections. Preserve explicit company, role title, dates, location, highlights, responsibilities, achievements, project details, and environment. Do not split one role unless clearly shown. Do not duplicate roles. Do not confuse role title with company name.
8. **CURRENT TITLE & NAME NORMALIZATION:** Preserve `basics.current_title` as close as possible to the resume header wording. Extract the real display name if explicit; if not, and the email clearly contains a safe `firstname.lastname` pattern, normalize it into a human-readable name. Do not invent beyond that.
9. **CANONICAL SECTION ROUTING:** Core content must go only into its canonical sections: `basics`, `summary`, `skills`, `experience`, `education`, `certifications`, `languages`. Degrees must go to `education`, certifications to `certifications`, and language items with proficiency/test details to `languages`.
10. **OTHER_SECTIONS ONLY:** Any remaining non-core content must go only into `other_sections`. Do not create, use, or reference `custom_sections`.
11. **CLEAN OUTPUT:** Use only empty strings `""` or arrays `[]` for missing values. Never output `None`, `null`, or placeholders. Keep wording faithful to the source but readable in professional US English. Avoid accidental ALL CAPS except for true acronyms or proper names.

**FINAL CHECK:**
- all human-readable text is in US English
- no explicit facts were lost
- no unsupported facts were invented
- no `None` or `null` appears
- skills from Top Skills / bullets / environment / responsibilities were not missed
- dates are preserved exactly as supported by the source
- no core content leaked into `other_sections`
- all remaining non-core content is preserved in `other_sections`
""",

    "prompt_qa": """Act as a strict QA Auditor for a recruitment agency. Compare the original attached CV with this extracted JSON.
NOTE: The JSON is generated programmatically. It purposefully forces skill categorization, infers standard job titles if missing, and standardizes empty fields to "". Do NOT report these architectural features as hallucinations.

EXTRACTED JSON:
{json_str}

TASK: Find ONLY real data losses or hallucinations. 
You MUST end your response with a JSON block in this EXACT format:
```json
{"score": 95, "missing": ["Skill 1", "Missing Date"], "hallucinations": ["Fake Certification"]}
```
If perfect, reply with score 100 and empty arrays.""",

    "prompt_qa_docx": """Act as a strict QA Auditor. Verify that the generated DOCX document faithfully represents the structured CV JSON.
NOTE: DOCX templates impose fixed formatting and section order — do NOT report layout or ordering differences as issues.

CV JSON:
{json_str}

TASK: Compare the JSON data against the attached DOCX text. Find ONLY real data losses or distortions:
- Key fields present in JSON but missing from DOCX (name, title, experience entries, skills, education, etc.)
- Text truncated or corrupted during rendering
- Wrong values (e.g. wrong dates, wrong company name) that contradict the JSON

You MUST end your response with a JSON block in this EXACT format:
```json
{"score": 95, "missing": ["Field or data absent from DOCX"], "hallucinations": ["Value in DOCX that contradicts JSON"]}
```
If perfect, reply with score 100 and empty arrays.""",

    "prompt_autofix": """You are a Self-Healing AI. Your task is to fix the JSON extraction of a CV based on a QA Audit report.
CURRENT (BROKEN) JSON:
{current_json_str}

QA AUDIT REPORT (ERRORS TO FIX):
{qa_report_text}

INSTRUCTIONS:
1. Look at the original attached CV file/text.
2. Fix ONLY the missing data or hallucinations mentioned in the QA report.
3. Maintain the EXACT same JSON schema as the CURRENT JSON. 
4. Do NOT remove any existing correct data.
5. Return ONLY the repaired JSON object without markdown wrappers.""",

    "prompt_matcher": """Act as a Senior IT Recruiter. Evaluate the candidate against the Job Description.
CRITICAL: Carefully analyze their actual 'experience' (duration, context, tools used), not just the 'skills' list.
First, think step-by-step about their fit.
Then, return a JSON ARRAY containing EXACTLY ONE object.

JD: {jd_val}
CANDIDATE: {cand_data}

SCHEMA:
[{
  "id": {i},
  "name": "Candidate Name",
  "reasoning": "Brief explanation of your evaluation based on their real experience",
  "score": integer (0-100),
  "verdict": "Short summary of fit",
  "pros": "Key strengths",
  "missing_skills": "Gaps or weaknesses"
}]""",

    "prompt_modifier": """You are an Expert CV Editor. Your task is to modify the provided Candidate JSON based EXACTLY on the user's request.

USER REQUEST:
"{user_req}"

CRITICAL RULES:
1. You MUST return ONLY a valid JSON object. No markdown formatting blocks, no explanations, no chat.
2. The returned JSON MUST strictly adhere to the exact same schema as the input JSON. Do not remove mandatory keys.
3. SAFE DELETION: If the user asks to remove, delete, or hide certain data (e.g., emails, phone numbers, specific jobs), DO NOT remove the keys from the JSON. Instead, set their values to an empty string `""` or an empty array `[]`.
4. AGGRESSIVE SHORTENING: If the user asks to shorten or reduce the CV (e.g., to 1 page), you MUST aggressively summarize: keep only the most recent/relevant jobs, limit achievements to 2-3 bullet points per job, and clear out secondary courses/hobbies by setting their values to empty strings/arrays.
5. Do not invent new work experience or skills unless the user asks you to infer or summarize existing ones.
6. LANGUAGE STRICTNESS: ALWAYS output the modified JSON content in professional US English.

INPUT JSON:
{input_json_str}""",

    "prompt_github": """You are an Expert Tech Recruiter. Convert this candidate's GitHub data into our STRICT JSON CV SCHEMA.
RULES:
1. Map GitHub 'name' (or 'login' if name is null) to `basics.name`.
2. Map 'location', 'email', 'blog', 'html_url' into the `basics` section.
3. Infer a professional `basics.current_title` based on their top languages.
4. Convert 'recent_repos' into the `projects` array.
5. Extract all programming languages used into the `skills` object.
6. Write a professional 2-3 sentence `summary.bullet_points` assessing their code footprint based on the "Code Quality Over Vanity" principle.
7. Leave `experience` and `education` as empty arrays `[]`.

JSON SCHEMA:
{prompt_schema_only}

GitHub API Data: {gh_full_data}""",

    "prompt_xray": """Act as an Expert Tech Recruiter.
Based on the following request, generate 3-5 advanced Google X-Ray Boolean search queries to find candidate profiles.
Target platforms: LinkedIn, GitHub, or general web.
Return ONLY a valid JSON array of objects, without Markdown formatting.
Schema:
[
  {"platform": "LinkedIn", "description": "Broad search for mid-level", "query": "site:linkedin.com/in (\\"Python\\" OR \\"Django\\") AND \\"AWS\\""}
]
Request: {user_input}""",

    "prompt_anonymize": """Act as a CV writer. Convert company names to generic industry descriptions (e.g., 'Large FinTech Company', 'Global E-commerce Enterprise'). Return JSON: {"Original": "Description"}.\nCompanies: {companies_json}"""
}


CURRENT_PROMPT_MASTER_VERSION = 2

DEFAULT_CONFIG = {
    "api_key": "", "github_token": "", "workspace_path": DEFAULT_WORKSPACE,
    "import_mode": "qa", "anon_cut_name": True, "anon_remove_creds": True,
    "anon_mask_companies": True, "keep_initial_current_title": False,
    "show_col_file": True, "show_col_company": True, "show_col_comments": True, "show_col_score": True,
    "show_xray_tab": False, "show_github_tab": False, "show_matcher_tab": False,
    "show_modify_tab": False, "show_qa_tab": False,
    "active_template": "quantori_classic.docx", 
    "json_naming_template": "CV_FirstName_LastName.json",
    "export_naming_template": "CV_FirstName_LastName.docx",
    "naming_template": "CV FirstName FirstLetter (CV_Alexei_L.docx)",
    "ui_theme": "Light", "last_jd": "", "last_sourcing_query": "", "last_modifier_query": "",
    "generate_docx_on_import": True,
    "prompt_master_version": CURRENT_PROMPT_MASTER_VERSION, "active_prompt_version": CURRENT_PROMPT_MASTER_VERSION,
    "prompt_master_user_edited": False, "_prompt_master_upgrade_warning": False,
    "last_qa_sample_size": "All available", "qa_compare_mode": "full_pipeline", "last_miner_keywords": "", "last_miner_location": "", "last_miner_stars": "100",
    "total_in_tokens": 0, "total_out_tokens": 0, "total_spent_usd": 0.0
}
# Merge default prompts into the main config
DEFAULT_CONFIG.update(DEFAULT_PROMPTS)

def _initial_master_prompts_registry():
    return {
        "active_version": CURRENT_PROMPT_MASTER_VERSION,
        "versions": {
            str(CURRENT_PROMPT_MASTER_VERSION): {
                "version": CURRENT_PROMPT_MASTER_VERSION,
                "title": "Current production baseline",
                "status": "active",
                "based_on": None,
                "notes": "Bootstrapped from code defaults.",
                "prompt_text": DEFAULT_PROMPTS["prompt_master_inst"],
            }
        }
    }


def load_master_prompts_registry():
    if os.path.exists(MASTER_PROMPTS_FILE):
        try:
            with open(MASTER_PROMPTS_FILE, 'r', encoding='utf-8') as f:
                reg = json.load(f)
                if isinstance(reg, dict):
                    return reg
        except Exception:
            pass
    return _initial_master_prompts_registry()


def save_master_prompts_registry(registry):
    with open(MASTER_PROMPTS_FILE, 'w', encoding='utf-8') as f:
        json.dump(registry, f, indent=2, ensure_ascii=False)


def ensure_master_prompts_registry():
    reg = load_master_prompts_registry()
    if not isinstance(reg.get("versions"), dict):
        reg["versions"] = {}
    key = str(CURRENT_PROMPT_MASTER_VERSION)
    if key not in reg["versions"]:
        reg["versions"][key] = {
            "version": CURRENT_PROMPT_MASTER_VERSION,
            "title": "Current production baseline",
            "status": "active" if reg.get("active_version", CURRENT_PROMPT_MASTER_VERSION) == CURRENT_PROMPT_MASTER_VERSION else "baseline",
            "based_on": None,
            "notes": "Bootstrapped from code defaults.",
            "prompt_text": DEFAULT_PROMPTS["prompt_master_inst"],
        }
    if not reg.get("active_version"):
        reg["active_version"] = CURRENT_PROMPT_MASTER_VERSION
    save_master_prompts_registry(reg)
    return reg


def get_master_prompt_entry(version=None, registry=None):
    reg = registry or ensure_master_prompts_registry()
    if version is None:
        version = reg.get("active_version", CURRENT_PROMPT_MASTER_VERSION)
    return (reg.get("versions") or {}).get(str(version))


def get_master_prompt_text(version=None, registry=None):
    entry = get_master_prompt_entry(version, registry=registry)
    if entry and entry.get("prompt_text"):
        return entry.get("prompt_text")
    return DEFAULT_PROMPTS["prompt_master_inst"]


def get_master_prompt_versions(registry=None):
    reg = registry or ensure_master_prompts_registry()
    return sorted([int(k) for k in (reg.get("versions") or {}).keys()], reverse=True)


def set_active_master_prompt_version(version, registry=None):
    reg = registry or ensure_master_prompts_registry()
    version = int(version)
    key = str(version)
    if key not in reg.get("versions", {}):
        raise KeyError(f"Prompt version v{version} not found")
    reg["active_version"] = version
    for k, v in reg.get("versions", {}).items():
        if isinstance(v, dict):
            v["status"] = "active" if int(k) == version else (v.get("status") if v.get("status") != "active" else "archived")
    save_master_prompts_registry(reg)
    return reg


def save_master_prompt_version(prompt_text, title="Prompt Editor save", notes="", based_on=None, make_active=True, status="experimental", registry=None):
    reg = registry or ensure_master_prompts_registry()
    versions = reg.setdefault("versions", {})
    existing = [int(k) for k in versions.keys()] or [CURRENT_PROMPT_MASTER_VERSION]
    new_ver = max(existing) + 1
    versions[str(new_ver)] = {
        "version": new_ver,
        "title": title,
        "status": status,
        "based_on": int(based_on) if based_on not in (None, "") else None,
        "notes": notes,
        "prompt_text": prompt_text,
    }
    if make_active:
        reg["active_version"] = new_ver
        for k, v in versions.items():
            if isinstance(v, dict):
                v["status"] = "active" if int(k) == new_ver else (v.get("status") if v.get("status") != "active" else "archived")
    save_master_prompts_registry(reg)
    return new_ver, reg


def load_config():
    ensure_master_prompts_registry()
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
                cfg = copy.deepcopy(DEFAULT_CONFIG)
                cfg.update(loaded)
                return cfg
        except Exception:
            pass
    cfg = copy.deepcopy(DEFAULT_CONFIG)
    save_config(cfg)
    return cfg


def save_config(cfg):
    with open(SETTINGS_FILE, 'w', encoding='utf-8') as f:
        json.dump(cfg, f, indent=4, ensure_ascii=False)

def init_workspace_folders(base_dir):
    if not os.path.exists(base_dir): os.makedirs(base_dir)
    folders = {
        "SOURCE": os.path.join(base_dir, 'source'),
        "JSON": os.path.join(base_dir, 'jsons'),
        "OUTPUT": os.path.join(base_dir, 'docxs'),
        "BLIND": os.path.join(base_dir, 'docxs_a'),
        "MODIFIED": os.path.join(base_dir, 'docxs_modified'),
        "REPORTS": os.path.join(base_dir, 'reports'),
        "TEMPLATES": os.path.join(base_dir, 'templates')
    }
    for folder in folders.values():
        if not os.path.exists(folder): os.makedirs(folder)
        
    target_template = os.path.join(folders["TEMPLATES"], "quantori_classic.docx")
    bundled_template = get_resource_path("quantori_classic.docx")
    if os.path.exists(bundled_template):
        # Always overwrite workspace template with the bundled one (template updates must win).
        # Keep a lightweight backup of the previous template if it already exists.
        try:
            if os.path.exists(target_template):
                bak = os.path.join(folders["TEMPLATES"], "quantori_classic.prev.docx")
                try:
                    shutil.copy2(target_template, bak)
                except Exception:
                    pass
            shutil.copy2(bundled_template, target_template)
        except Exception:
            pass
    return folders

def open_folder(path):
    if not os.path.exists(path): os.makedirs(path, exist_ok=True)
    if platform.system() == "Windows": os.startfile(path)
    elif platform.system() == "Darwin": subprocess.Popen(["open", path])
    else: subprocess.Popen(["xdg-open", path])

# ==========================================
# 3. JSON SANITIZATION
# ==========================================

def fix_company_name_artifacts(data: dict) -> dict:
    """
    Fix common parsing artifacts where a section label is mistakenly captured as company_name.
    Example: company_name == "Project" / "Accomplishments" / "Environment".
    Lossless approach: we do NOT delete content, only clear the incorrect company_name value.
    """
    if not isinstance(data, dict):
        return data

    BAD = {
        "project", "projects", "project:", "projects:",
        "accomplishment", "accomplishments", "accomplishments:",
        "environment", "environment:", "responsibilities", "responsibilities:",
        "role", "role:", "position", "position:", "company", "company:",
        "time period", "time period:", "timeperiod", "timeperiod:",
    }

    def _is_bad(s: str) -> bool:
        if not isinstance(s, str):
            return False
        t = s.strip().lower()
        if not t:
            return False
        # exact matches
        if t in BAD:
            return True
        # common patterns like "Project" surrounded by punctuation
        if t.rstrip(".") in BAD:
            return True
        return False

    # v1: experience[]
    exp = data.get("experience")
    if isinstance(exp, list):
        for job in exp:
            if isinstance(job, dict):
                cn = job.get("company_name")
                if isinstance(cn, str) and _is_bad(cn):
                    job["company_name"] = ""

    # v2: work_experience[]
    wexp = data.get("work_experience")
    if isinstance(wexp, list):
        for job in wexp:
            if isinstance(job, dict):
                cn = job.get("company_name")
                if isinstance(cn, str) and _is_bad(cn):
                    job["company_name"] = ""

    return data


def _collect_raw_text_for_languages(data: dict) -> str:
    try:
        raw = data.get("raw")
        if isinstance(raw, str):
            return raw
        if isinstance(raw, dict):
            parts = []
            for v in raw.values():
                if isinstance(v, str):
                    parts.append(v)
                elif isinstance(v, list):
                    parts.extend([x for x in v if isinstance(x, str)])
            return "\n".join(parts)
        if isinstance(raw, list):
            return "\n".join([x for x in raw if isinstance(x, str)])
    except Exception:
        pass
    return ""

def _short_lang_level(s: str) -> str:
    if not isinstance(s, str):
        return ""
    t = s.strip()
    if not t:
        return ""
    # Prefer CEFR token if present
    m = re.search(r"\b([ABC][12])\b", t, flags=re.I)
    if m:
        return m.group(1).upper()
    # Normalize native variants
    if re.search(r"\bnative\b", t, flags=re.I):
        return "Native"
    return t

def ensure_native_languages(data: dict) -> dict:
    """If raw text contains 'NATIVE <LANG>' and it's not in languages[], add it losslessly."""
    if not isinstance(data, dict):
        return data
    raw_text = _collect_raw_text_for_languages(data)
    if not raw_text:
        return data
    langs = data.get("languages")
    if not isinstance(langs, list):
        return data

    existing = set()
    for it in langs:
        if isinstance(it, dict):
            name = (it.get("language") or "").strip().lower()
            if name:
                existing.add(name)
        elif isinstance(it, str):
            existing.add(it.strip().lower())

    # capture patterns like 'NATIVE RUSSIAN' or 'Native Russian'
    for m in re.finditer(r"\bNATIVE\s+([A-Z][A-Z]+|[A-Z][a-z]+)\b", raw_text, flags=re.I):
        lang = m.group(1).strip()
        if not lang:
            continue
        # Title-case if all caps
        if lang.isupper():
            lang = lang.title()
        key = lang.lower()
        if key in existing:
            continue
        if not _is_probably_tech_language(lang):
            langs.append({"language": lang, "proficiency": "Native", "level": "Native", "details": "Native"})
        existing.add(key)

    data["languages"] = langs
    return data



def _strip_leading_list_marker_text(s: str) -> str:
    """Remove leading bullet/list markers after optional indentation.

    This is intentionally conservative: it only trims marker glyphs at the
    start of a string (after whitespace). Inline dashes/bullets inside normal
    text are preserved.
    """
    if not isinstance(s, str):
        return s
    # Normalize NBSP that often appears after copied bullets.
    s = s.replace(" ", " ")
    # Repeated marker sequences like "• • AWS" / "-- item" / "— item"
    return re.sub(r'^\s*(?:[•●■▪▫◦‣⁃∙·*\-–—]+\s*)+', '', s)


def _strip_leading_list_markers_deep(obj):
    """Recursively remove leading list markers from all incoming text fields."""
    if isinstance(obj, str):
        return _strip_leading_list_marker_text(obj)
    if isinstance(obj, list):
        return [_strip_leading_list_markers_deep(x) for x in obj]
    if isinstance(obj, dict):
        return {k: _strip_leading_list_markers_deep(v) for k, v in obj.items()}
    return obj


def _normalize_human_language(name: str) -> str:
    """Normalize language name for whitelist matching."""
    if not isinstance(name, str):
        return ""
    s = name.strip()
    if not s:
        return ""
    # remove leading bullets and punctuation
    s = re.sub(r"^[\s•●■▪▫◦‣⁃∙·–—\-]+", "", s).strip()
    # drop CEFR / level in parentheses or after dash, e.g. "English (C1)" / "English - C1"
    s = re.sub(r"\([^)]*\)", "", s).strip()
    s = re.sub(r"\s*[-–—]\s*(A1|A2|B1|B2|C1|C2|Native|Fluent|Advanced|Intermediate|Beginner)\b.*$", "", s, flags=re.I).strip()
    # drop trailing commas/colons
    s = s.rstrip(" ,;:")
    return s

_HUMAN_LANG_ALIASES = {
    # abbreviations
    "en": "english",
    "eng": "english",
    "ru": "russian",
    "rus": "russian",
    "de": "german",
    "ger": "german",
    "fr": "french",
    "es": "spanish",
    "spa": "spanish",
    "pt": "portuguese",
    "por": "portuguese",
    "it": "italian",
    "nl": "dutch",
    "uk": "ukrainian",
    "ua": "ukrainian",
    "hy": "armenian",
    "ka": "georgian",
    "zh": "chinese",
    "ja": "japanese",
    "jp": "japanese",
    "ko": "korean",
    "ar": "arabic",
    "fa": "persian",
    "he": "hebrew",
}

_HUMAN_LANG_WHITELIST = {
    # top common
    "english","russian","german","french","spanish","portuguese","italian","dutch","ukrainian","polish","turkish",
    "armenian","georgian","azerbaijani","kazakh","uzbek","tajik","kyrgyz","belarusian","moldovan","romanian","bulgarian",
    "greek","serbian","croatian","bosnian","slovenian","slovak","czech","hungarian","albanian","macedonian","montenegrin",
    "swedish","norwegian","danish","finnish","icelandic","estonian","latvian","lithuanian",
    "chinese","mandarin","cantonese","japanese","korean","vietnamese","thai","indonesian","malay","filipino","tagalog",
    "hindi","urdu","bengali","punjabi","tamil","telugu","marathi","gujarati","kannada","malayalam","sinhala","nepali",
    "arabic","persian","farsi","hebrew","kurdish",
    "swahili","afrikaans","zulu","xhosa","amharic",
    "latin","irish","scottish gaelic","welsh",
    # variants
    "persion","french","spanish","portuguese","brazilian portuguese","brazilian","português",
}

def _is_human_language(name: str) -> bool:
    """Strict whitelist-based classifier for human languages."""
    n0 = _normalize_human_language(name)
    if not n0:
        return False
    n = n0.lower()
    # normalize common alias tokens (two/three-letter codes)
    if n in _HUMAN_LANG_ALIASES:
        n = _HUMAN_LANG_ALIASES[n]
    # normalize 'farsi' => persian
    if n == "farsi":
        n = "persian"
    # normalize multiple spaces
    n = re.sub(r"\s+", " ", n).strip()
    return n in _HUMAN_LANG_WHITELIST

def sync_languages_to_skills(data: dict) -> dict:
    """Render human Languages into skills['Languages'] with levels, losslessly.
    Strict: only whitelist human languages go into this list.
    Non-matching tokens are preserved in extras as 'Languages(unclassified): ...' (lossless).

    Additionally, remove duplicates like 'English' + 'English (C1)' by keeping the richer version.
    """
    if not isinstance(data, dict):
        return data

    langs = data.get("languages")
    if not isinstance(langs, list) or not langs:
        return data

    def _base_name(s: str) -> str:
        s = _normalize_human_language(s)
        s = re.sub(r"\s+", " ", s).strip().lower()
        if s in _HUMAN_LANG_ALIASES:
            s = _HUMAN_LANG_ALIASES[s]
        if s == "farsi":
            s = "persian"
        return s

    rendered = []
    unclassified = []

    for it in langs:
        if isinstance(it, dict):
            name = (it.get("language") or "").strip()
            if not name:
                continue
            if not _is_human_language(name):
                unclassified.append(name)
                continue
            lvl = _short_lang_level(it.get("level") or it.get("proficiency") or it.get("details") or "")
            rendered.append(f"{name} ({lvl})" if lvl else name)
        elif isinstance(it, str) and it.strip():
            name = it.strip()
            if not _is_human_language(name):
                unclassified.append(name)
                continue
            rendered.append(name)

    # Merge previous Languages entries (only human ones), but don't re-introduce poorer duplicates.
    skills = data.get("skills")
    if not isinstance(skills, dict):
        skills = {}

    prev = skills.get("Languages")
    if isinstance(prev, list):
        for x in prev:
            if isinstance(x, str) and x.strip():
                nm = x.strip()
                if _is_human_language(nm):
                    rendered.append(nm)

    # Choose best variant per base language name, preserving first-seen order of bases.
    best_by_base = {}
    order = []
    for x in rendered:
        base = _base_name(x)
        if not base:
            continue
        if base not in best_by_base:
            best_by_base[base] = x
            order.append(base)
        else:
            cur = best_by_base[base]
            # Prefer strings that include a level "(C1)" / "(Native)" or are longer.
            def score(v: str) -> int:
                s = 0
                if re.search(r"\([ABC][12]\)", v):
                    s += 5
                if re.search(r"\(Native\)", v, flags=re.I):
                    s += 4
                if "(" in v and ")" in v:
                    s += 2
                s += min(3, len(v)//10)
                return s
            if score(x) > score(cur):
                best_by_base[base] = x

    rendered2 = [best_by_base[b] for b in order if b in best_by_base]

    skills["Languages"] = rendered2
    data["skills"] = skills

    # Lossless: stash unclassified tokens so they are not lost
    if unclassified:
        extras = data.get("extras")
        if not isinstance(extras, list):
            extras = []
        for tok in unclassified:
            line = f"Languages(unclassified): {tok}"
            if line not in extras:
                extras.append(line)
        data["extras"] = extras

    return data



_MONTH_NAMES = ["january","february","march","april","may","june",
                "july","august","september","october","november","december"]

def _is_future_date(s):
    """Return True if date string represents a month/year clearly in the future."""
    if not isinstance(s, str): return False
    low = s.strip().lower()
    if not low or low == 'present': return False
    m = re.search(r'\b(20\d\d|19\d\d)\b', s)
    if not m: return False
    year = int(m.group(1))
    today = datetime.date.today()
    if year > today.year: return True
    if year == today.year:
        for i, name in enumerate(_MONTH_NAMES):
            if name in low:
                return (i + 1) > today.month
    return False


def sanitize_json(data):
    if not isinstance(data, dict): data = {}
    data = _strip_leading_list_markers_deep(data)
    
    backup_keys = {k: data.get(k) for k in ['qa_audit', 'match_analysis', '_source_filename', '_source_hash', 'import_date', '_comment']}

    if not isinstance(data.get('basics'), dict): data['basics'] = {}
    bad_values = ['unavailable', 'n/a', 'none', 'null', 'not provided', 'unknown', 'undisclosed']
    
    for key in ['name', 'location', 'objective', 'current_company']:
            val = data['basics'].get(key)
            if isinstance(val, str):
                cleaned_val = val.strip().replace('\n', ' ').replace('\r', '')
                if key == 'name' and cleaned_val:
                    if cleaned_val.isupper() or cleaned_val.islower(): cleaned_val = cleaned_val.title()
                data['basics'][key] = "" if cleaned_val.lower() in bad_values else cleaned_val
            else: 
                data['basics'][key] = ", ".join(map(str, val)) if isinstance(val, list) else (str(val) if val else "")

    raw_title = data['basics'].get('current_title', '')
    if isinstance(raw_title, str):
            raw_title = raw_title.replace('\n', ' ').replace('\r', '').strip()
            if raw_title.lower() in bad_values: raw_title = ""
    else: raw_title = str(raw_title) if raw_title else ""
            
    data['basics']['current_title_original'] = raw_title
            
    clean_title = raw_title
    if clean_title:
            clean_title = re.split(r'\s*\|\s*|\s*-\s*|\s*,\s*|\s+at\s+|\s+@\s+', clean_title, flags=re.IGNORECASE)[0].strip()
            if clean_title.isupper(): clean_title = clean_title.title()
            if len(clean_title) > 100 or len(clean_title) < 2: clean_title = ""

    if not clean_title and data.get('experience'):
            fb = str(data['experience'][0].get('role', '')).replace('\n', ' ').strip()
            if fb: clean_title = fb
            
    data['basics']['current_title'] = clean_title

    if not isinstance(data['basics'].get('contacts'), dict): data['basics']['contacts'] = {}
    clean_contacts = {}
    for k, v in data['basics']['contacts'].items():
        s_val = ", ".join(map(str, v)) if isinstance(v, list) else str(v)
        if s_val and s_val.lower().strip() not in bad_values: clean_contacts[k] = s_val
    data['basics']['contacts'] = clean_contacts
    
    if not isinstance(data['basics'].get('links'), list): data['basics']['links'] = []
    data['basics']['links'] = [str(l) for l in data['basics']['links'] if l and str(l).lower().strip() not in bad_values]

    if 'skills' not in data or not isinstance(data['skills'], dict): data['skills'] = {}
    clean_skills = {}
    for k, v in data['skills'].items():
        # Normalize schema-leaked key names: "technical_skills" -> "Technical Skills"
        clean_k = k.replace('_', ' ').strip().title() if '_' in k else k
        if isinstance(v, str): clean_skills[clean_k] = [v]
        elif isinstance(v, list): clean_skills[clean_k] = [str(i) for i in v if i]
    data['skills'] = clean_skills

    if not isinstance(data.get('experience'), list): data['experience'] = []
    for job in data['experience']:
        pd = job.get('project_description')
        if isinstance(pd, list): pd = " ".join(map(str, pd))
        job['project_description'] = "" if (isinstance(pd, str) and pd.lower().strip() in bad_values) else (pd if isinstance(pd, str) else "")
            
        loc = job.get('location')
        job['location'] = "" if (isinstance(loc, str) and loc.lower().strip() in bad_values) else (loc if isinstance(loc, str) else "")
        
        hl = job.get('highlights')
        job['highlights'] = [hl] if isinstance(hl, str) else (hl if isinstance(hl, list) else [])
        
        env = job.get('environment')
        job['environment'] = [env] if isinstance(env, str) else (env if isinstance(env, list) else [])
        
        if 'dates' in job:
            for d_key in ['start', 'end']:
                val = job['dates'].get(d_key)
                job['dates'][d_key] = "" if (isinstance(val, str) and val.lower().strip() in bad_values) else (val if isinstance(val, str) else "")
            # Future end date → replace with "Present" (LLM hallucination guard)
            end_val = job['dates'].get('end', '')
            if _is_future_date(end_val):
                job['dates']['end'] = 'Present'
        
        c_val = job.get('company_name')
        job['company_name'] = "" if (isinstance(c_val, str) and c_val.lower().strip() in bad_values) else (c_val if isinstance(c_val, str) else "")

    # Deduplicate experience entries by (company_name, role, start_date)
    seen_exp = set()
    clean_exp = []
    for job in data['experience']:
        key = (
            str(job.get('company_name', '')).strip().lower(),
            str(job.get('role', '')).strip().lower(),
            str((job.get('dates') or {}).get('start', '')).strip().lower(),
        )
        if key in seen_exp:
            continue
        seen_exp.add(key)
        clean_exp.append(job)
    data['experience'] = clean_exp

    if not isinstance(data.get('projects'), list): data['projects'] = []
    clean_projects = []
    for p in data['projects']:
        if not isinstance(p, dict): p = {}
        for k in ['title', 'description', 'link']:
            val = p.get(k)
            p[k] = "" if (isinstance(val, str) and val.lower().strip() in bad_values) else (val if isinstance(val, str) else "")
            
        if 'tech_stack' not in p or not isinstance(p['tech_stack'], list): p['tech_stack'] = []
        p['tech_stack'] = [str(t) for t in p['tech_stack'] if t and str(t).lower().strip() not in bad_values]
        if p.get('title') or p.get('description'): clean_projects.append(p)
    data['projects'] = clean_projects
        
    for k in ['certifications', 'publications', 'courses']:
        if k not in data or not isinstance(data[k], list): data[k] = []
        else: data[k] = [str(x) for x in data[k] if x and str(x).lower().strip() not in bad_values]

    if 'education' in data and isinstance(data['education'], list):
        for edu in data['education']:
            for e_key in ['institution', 'degree', 'year', 'details']:
                val = edu.get(e_key)
                if val is None:
                    edu[e_key] = ""
                elif isinstance(val, str) and val.lower().strip() in bad_values:
                    edu[e_key] = ""
        # Deduplicate education entries by (institution, degree, year)
        seen_edu = set()
        clean_edu = []
        for edu in data['education']:
            key = (
                str(edu.get('institution', '')).strip().lower(),
                str(edu.get('degree', '')).strip().lower(),
                str(edu.get('year', '')).strip().lower(),
            )
            if key in seen_edu:
                continue
            seen_edu.add(key)
            clean_edu.append(edu)
        data['education'] = clean_edu

    if 'volunteering' not in data or not isinstance(data['volunteering'], list): data['volunteering'] = []
    for v in data['volunteering']:
        if not isinstance(v, dict): continue
        for k in ['organization', 'role']:
            val = v.get(k)
            v[k] = "" if (isinstance(val, str) and val.lower().strip() in bad_values) else (str(val) if val else "")
        if 'highlights' not in v or not isinstance(v['highlights'], list): v['highlights'] = []
        else: v['highlights'] = [str(h) for h in v['highlights'] if h]

    data = normalize_languages_lossless(data)
    if 'languages' not in data or not isinstance(data['languages'], list): data['languages'] = []
    for l in data['languages']:
        if not isinstance(l, dict): continue
        # keep extra fields (level/details), but sanitize placeholders
        for k in ['language','proficiency','level','details']:
            if k not in l: continue
            val = l.get(k)
            l[k] = "" if (isinstance(val, str) and val.lower().strip() in bad_values) else (str(val) if val else "")

    # Merge any legacy custom_sections into other_sections, then remove the legacy field
    # so other_sections remains the only non-core bucket in the final JSON.
    if not isinstance(data.get('other_sections'), list):
        data['other_sections'] = []

    def _normalize_other_section(sec, *, title_keys=("title", "section_title")):
        if not isinstance(sec, dict):
            return None
        title = ""
        for key in title_keys:
            raw_title = sec.get(key, "")
            if isinstance(raw_title, str):
                raw_title = raw_title.strip()
            elif raw_title:
                raw_title = str(raw_title).strip()
            else:
                raw_title = ""
            if raw_title:
                title = raw_title
                break
        if isinstance(title, str) and title.lower().strip() in bad_values:
            title = ""

        raw_items = sec.get('items', [])
        if isinstance(raw_items, list):
            items = [str(x).strip() for x in raw_items if x and str(x).strip() and str(x).lower().strip() not in bad_values]
        elif raw_items and str(raw_items).strip() and str(raw_items).lower().strip() not in bad_values:
            items = [str(raw_items).strip()]
        else:
            items = []

        if not title and not items:
            return None
        return {"title": title, "items": items}

    merged_other = []
    for sec in data.get('other_sections', []):
        norm = _normalize_other_section(sec, title_keys=("title", "section_title"))
        if norm:
            merged_other.append(norm)
    # Migrate legacy top-level non-core fields into other_sections
    # so other_sections becomes the only canonical non-core bucket.
    def _project_lines(projects):
        lines = []
        if not isinstance(projects, list):
            return lines
        for p in projects:
            if isinstance(p, dict):
                title = str(p.get("title") or p.get("name") or "").strip()
                desc = str(p.get("description") or "").strip()
                link = str(p.get("link") or p.get("url") or "").strip()
                line = title or ""
                if link:
                    line = (line + " — " + link).strip(" —")
                if desc:
                    line = (line + " — " + desc).strip(" —")
                if line:
                    lines.append(line)
            elif isinstance(p, str) and p.strip():
                lines.append(p.strip())
        return lines

    legacy_map = [
        ("projects", "Projects", _project_lines),
        ("courses", "Courses", lambda v: [str(x).strip() for x in v if str(x).strip()] if isinstance(v, list) else []),
        ("publications", "Publications", lambda v: [str(x).strip() for x in v if str(x).strip()] if isinstance(v, list) else []),
        ("volunteering", "Volunteering", lambda v: [
            " — ".join([
                str(it.get("organization", "")).strip(),
                str(it.get("role", "")).strip()
            ]).strip(" —")
            for it in v if isinstance(it, dict) and (
                str(it.get("organization", "")).strip() or str(it.get("role", "")).strip()
            )
        ] if isinstance(v, list) else []),
        ("extras", "Other", lambda v: [str(x).strip() for x in v if str(x).strip()] if isinstance(v, list) else []),
        ("other", "Other", lambda v: [str(x).strip() for x in v if str(x).strip()] if isinstance(v, list) else []),
    ]

    for key, title, fn in legacy_map:
        lines = fn(data.get(key))
        if lines:
            merged_other.append({"title": title, "items": lines})
    legacy_custom = data.get('custom_sections', [])
    if isinstance(legacy_custom, list):
        for sec in legacy_custom:
            norm = _normalize_other_section(sec, title_keys=("section_title", "title"))
            if norm:
                merged_other.append(norm)

    # De-duplicate by (title, items) while preserving order.
    seen_sections = set()
    clean_other = []
    for sec in merged_other:
        sig = (sec.get('title', '').strip().casefold(), tuple(i.casefold() for i in sec.get('items', [])))
        if sig in seen_sections:
            continue
        seen_sections.add(sig)
        clean_other.append(sec)

    # Remove duplicated core sections from other_sections if canonical core is already filled
    filtered_other = []
    for sec in clean_other:
        title = (sec.get("title", "") or "").strip().casefold()
        if title == "languages" and data.get("languages"):
            continue
        if title == "certifications" and data.get("certifications"):
            continue
        if title == "education" and data.get("education"):
            continue
        filtered_other.append(sec)

    data['other_sections'] = filtered_other

    # Remove legacy non-core containers from final JSON
    for legacy_key in [
        'custom_sections',
        'projects',
        'courses',
        'publications',
        'volunteering',
        'extras',
        'other',
    ]:
        data.pop(legacy_key, None)

    # Remove obvious section-label artifacts (e.g., company_name == 'Project')
    data = fix_company_name_artifacts(data)

    for k, v in backup_keys.items():
        if v is not None: data[k] = v

    return data



def normalize_languages_lossless(data: dict) -> dict:
    """
    Lossless normalization for languages:
    - Accept list items as dict or string
    - Preserve extra fields (level/details/notes)
    - Extract CEFR-like level from proficiency/details when possible
    """
    if not isinstance(data, dict):
        return data
    langs = data.get("languages")
    if langs is None:
        return data
    if not isinstance(langs, list):
        langs = [langs]
    out = []
    # CEFR / common level tokens
    level_re = re.compile(r"\b(A1|A2|B1|B2|C1|C2)\b", re.I)
    native_re = re.compile(r"\b(native|mother\s*tongue|fluent|bilingual)\b", re.I)

    for item in langs:
        if isinstance(item, str):
            s = item.strip()
            if not s:
                continue
            # Try split "Language - Level/Details"
            # Examples: "English — C1 (Advanced)", "Russian (Native)"
            lang = s
            prof = ""
            # split on dash-like separators
            parts = re.split(r"\s*(?:[-–—:|]|,)\s*", s, maxsplit=1)
            if len(parts) == 2 and parts[0] and parts[1]:
                lang = parts[0].strip()
                prof = parts[1].strip()
            d = {"language": lang, "proficiency": prof}
            out.append(d)
            continue

        if isinstance(item, dict):
            # keep everything, but ensure canonical keys exist
            d = dict(item)
            lang = str(d.get("language") or d.get("name") or "").strip()
            prof = str(d.get("proficiency") or d.get("level") or "").strip()
            details = str(d.get("details") or d.get("notes") or "").strip()

            # If language is missing but we have a single key dict like {"English": "C1"} treat it
            if not lang and len(d) == 1:
                k = next(iter(d.keys()))
                v = d[k]
                if isinstance(k, str):
                    lang = k.strip()
                    prof = str(v).strip() if v is not None else prof

            # Extract level from prof/details
            level = str(d.get("level") or "").strip()
            if not level:
                m = level_re.search(prof) or level_re.search(details)
                if m:
                    level = m.group(1).upper()
                else:
                    m2 = native_re.search(prof) or native_re.search(details)
                    if m2:
                        level = m2.group(1).title()

            # If details empty but prof contains more than just level, keep it
            if not details and prof and (len(prof) > 4 or "(" in prof or "/" in prof):
                details = prof

            d["language"] = lang
            d["proficiency"] = prof
            if level:
                d["level"] = level
            if details:
                d["details"] = details

            # Drop known bad placeholders but keep lossless extras
            if lang or prof or details or level:
                out.append(d)
            continue

    data["languages"] = out
    
    # Languages: add missing Native entries from raw and render levels into skills['Languages']
    data = ensure_native_languages(data)
    data = sync_languages_to_skills(data)

    return data

def _collect_raw_text(ctx: dict) -> str:
    parts = []
    raw = ctx.get("raw") or {}
    if isinstance(raw, dict):
        for k in ("source_text_chunks","unmapped_facts"):
            v = raw.get(k)
            if isinstance(v, list):
                parts.extend([x for x in v if isinstance(x,str)])
            elif isinstance(v, str):
                parts.append(v)
    for k in ("source_text","raw_text","text","_source_text"):
        v = ctx.get(k)
        if isinstance(v, str):
            parts.append(v)
    return "\n".join([p.strip() for p in parts if isinstance(p,str) and p.strip()])

def _format_docx_sections_for_llm(baseline: dict) -> str:
    """Format structured DOCX sections into labeled text for the LLM."""
    sections = baseline.get("sections", {})
    section_labels = {
        "preamble": "CONTACT INFO (name, location, email — NOT the job title)",
        "summary": "SUMMARY",
        "skills": "TECHNICAL SKILLS",
        "experience": "WORK EXPERIENCE",
        "education": "EDUCATION",
        "certifications": "CERTIFICATIONS",
        "languages": "LANGUAGES",
    }
    parts = []
    for key, label in section_labels.items():
        lines = sections.get(key, [])
        if lines:
            parts.append(f"=== {label} ===\n" + "\n".join(lines))
    for key, lines in sections.items():
        if key not in section_labels and lines:
            parts.append(f"=== {key.upper().replace('_', ' ')} ===\n" + "\n".join(lines))
    return "\n\n".join(parts)


def extract_text_from_docx(docx_path: str) -> str:
    """Robust DOCX text extraction for CVs (paragraphs + tables + headers/footers)."""
    try:
        from docx import Document
    except Exception:
        return ""
    try:
        doc = Document(docx_path)
    except Exception:
        return ""

    parts = []

    def add_paras(paras):
        for p in paras:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

    add_paras(doc.paragraphs)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                add_paras(cell.paragraphs)

    try:
        for sec in doc.sections:
            add_paras(sec.header.paragraphs)
            add_paras(sec.footer.paragraphs)
    except Exception:
        pass

    # dedupe consecutive identical lines
    cleaned = []
    prev = None
    for line in parts:
        if line == prev:
            continue
        cleaned.append(line)
        prev = line

    return "\n".join(cleaned).strip()

# ==========================================
# 4. LLM PROCESSING CORE
# ==========================================
def process_file_gemini(file_path, api_key, custom_instructions, task_state=None): 
    # 🧠 Concatenate editable instructions and the protected schema
    final_prompt = custom_instructions + f"\n\n**JSON SCHEMA:**\n{CV_JSON_SCHEMA}"
    
    client = genai.Client(api_key=api_key)

    if file_path.lower().endswith('.docx'):
        try:
            from source_baseline_extractor import extract_from_docx as _extract_from_docx
            baseline = _extract_from_docx(file_path)
            text = _format_docx_sections_for_llm(baseline)
        except Exception:
            text = extract_text_from_docx(file_path)
        if not text: raise ValueError("Empty DOCX")
        response = _retry_generate(client, MODEL_NAME, [final_prompt, text])
    else:
        mime = 'application/pdf' if file_path.lower().endswith('.pdf') else 'image/jpeg'
        import unicodedata as _ud
        _safe_name = _ud.normalize('NFKD', os.path.basename(file_path)).encode('ascii', 'ignore').decode('ascii') or "cv_file"
        sample = client.files.upload(file=file_path, config=genai_types.UploadFileConfig(mime_type=mime, display_name=_safe_name))
        while sample.state.name == "PROCESSING":
            if task_state and task_state.get("cancel"): return None, 0, 0, 0.0
            time.sleep(1)
            sample = client.files.get(name=sample.name)
        if task_state and task_state.get("cancel"): return None, 0, 0, 0.0
        response = _retry_generate(client, MODEL_NAME, [sample, final_prompt])
        
    text = response.text.replace('```json', '').replace('```', '').strip()
    
    in_tok = getattr(response.usage_metadata, 'prompt_token_count', 0)
    out_tok = getattr(response.usage_metadata, 'candidates_token_count', 0)
    cost = (in_tok / 1_000_000 * PRICE_1M_IN) + (out_tok / 1_000_000 * PRICE_1M_OUT)
    
    if not text: return None, in_tok, out_tok, cost
    data = json.loads(text)
    return sanitize_json(data), in_tok, out_tok, cost

def generate_docx_from_json(data, output_path, cfg):
    workspace = cfg.get("workspace_path", DEFAULT_WORKSPACE)
    template_dir = os.path.join(workspace, "templates")
    template_name = cfg.get("active_template", "quantori_classic.docx")
    template_path = os.path.join(template_dir, template_name)
    
    if not os.path.exists(template_path):
        fallback = os.path.join(template_dir, "quantori_classic.docx")
        if os.path.exists(fallback):
            template_path = fallback
        else:
            raise FileNotFoundError(f"Template not found! Please create: {template_path}")
        
    doc = DocxTemplate(template_path)
    context = _trim_strings_deep(copy.deepcopy(data))
    # Ensure required keys exist for Jinja templates (avoid StrictUndefined crashes)
    if not isinstance(context, dict):
        context = {}
    context.setdefault('basics', {})
    context.setdefault('summary', {})
    context.setdefault('skills', {})
    context.setdefault('experience', [])
    context.setdefault('education', [])
    context.setdefault('certifications', [])
    context.setdefault('projects', [])
    context.setdefault('extras', [])
    
    c_list = [str(v).strip() for k, v in context.get('basics', {}).get('contacts', {}).items() if str(v).strip()]
    if str(context.get('basics', {}).get('location', '')).strip(): 
        c_list.append(str(context['basics']['location']).strip())
    c_list.extend([str(x).strip() for x in context.get('basics', {}).get('links', []) if str(x).strip()])
    context['contact_line'] = " | ".join(c_list)
    
    if cfg.get('keep_initial_current_title', False):
        context['basics']['current_title'] = context.get('basics', {}).get('current_title_original', context.get('basics', {}).get('current_title', ''))
    # Optional: enrich contact_line from raw text if empty
    raw_text = _collect_raw_text(context)
    if not context.get('contact_line'):
        c2 = _extract_contacts_plus(raw_text)
        parts = []
        for k in ('email', 'phone', 'website', 'linkedin'):
            v = c2.get(k)
            if isinstance(v, str) and v.strip():
                parts.append(v.strip())
        loc = _extract_location_line(raw_text)
        if isinstance(loc, str) and loc.strip():
            parts.append(loc.strip())
        if parts:
            # de-dupe while preserving order
            context['contact_line'] = " | ".join(dict.fromkeys(parts))

    
    # Summary normalization: move bullet points into Summary (regulations have no "Key Highlights")
    context.setdefault("summary", {})
    if not isinstance(context["summary"], dict):
        context["summary"] = {}

    # v2 -> v1 mapping
    items = context["summary"].get("items") or []
    if items and not context["summary"].get("bullet_points"):
        context["summary"]["bullet_points"] = [x.strip() for x in items if isinstance(x, str) and x.strip()]

    # If still no bullets, derive from objective (best-effort)
    # Only when objective is the sole source of content — clear it afterwards to avoid
    # the template rendering both objective and summary with identical text.
    if not context["summary"].get("bullet_points"):
        obj = (context.get("basics") or {}).get("objective")
        if isinstance(obj, str) and obj.strip():
            parts = re.split(r"(?<=[.!?])\s+", obj.strip())
            bullets = [p.strip() for p in parts if p.strip()]
            # limit to 7
            context["summary"]["bullet_points"] = bullets[:7]
            context["basics"]["objective"] = ""


    # Build other_sections for any non-required content (goes to the end of CV under its own headings)
    # Required sections per regulations are rendered earlier (Summary/Technical Skills/Work Experience/Education/Certifications).
    other_sections = []

    def _as_lines(val):
        if val is None:
            return []
        if isinstance(val, str):
            return [val.strip()] if val.strip() else []
        if isinstance(val, list):
            out=[]
            for x in val:
                if isinstance(x, str) and x.strip():
                    out.append(x.strip())
                elif isinstance(x, dict):
                    # best-effort stringify common keys
                    txt = " — ".join([str(x.get(k,"")).strip() for k in ("title","name","language","institution","degree","description") if str(x.get(k,"")).strip()])
                    if txt:
                        out.append(txt)
            return out
        if isinstance(val, dict):
            # dict of category->list
            out=[]
            for k,v in val.items():
                if isinstance(v, list) and v:
                    items=[str(i).strip() for i in v if str(i).strip()]
                    if items:
                        out.append(f"{k}: " + ", ".join(items))
                elif isinstance(v, str) and v.strip():
                    out.append(f"{k}: {v.strip()}")
            return out
        return [str(val)]
    
#=========================================================
    # Candidate optional sections (canonical only)
    # Languages: render ONLY explicit human/spoken languages from the canonical
    # structured `languages` section. Never derive this section from skills["Languages"].
    def _as_lines(lang_items):
        out = []
        if not isinstance(lang_items, list):
            return out
        for it in lang_items:
            if isinstance(it, dict):
                lang = str(it.get("language") or "").strip()
                prof = str(it.get("proficiency") or it.get("level") or it.get("details") or "").strip()
                if lang and prof:
                    out.append(f"{lang} ({prof})")
                elif lang:
                    out.append(lang)
            elif isinstance(it, str) and it.strip():
                out.append(it.strip())
        return out

    lang_lines = _as_lines(context.get("languages"))
    if lang_lines:
        other_sections.append({"title": "Languages", "items": lang_lines})
        # Remove "Languages" key from skills dict to prevent double rendering
        # (once under Technical Skills, once in the dedicated Languages section).
        skills = context.get("skills")
        if isinstance(skills, dict):
            keys_to_remove = [k for k in skills if k.strip().lower() == "languages"]
            for k in keys_to_remove:
                del skills[k]

#===========================================================================            

    # Merge any pre-normalized other_sections from JSON/context so template uses a single
    # bucket for all non-core sections.
    existing_other = context.get("other_sections")
    if isinstance(existing_other, list):
        for sec in existing_other:
            if not isinstance(sec, dict):
                continue
            title = str(sec.get("title", "")).strip()
            items = sec.get("items", [])
            if not isinstance(items, list):
                items = [items] if items else []
            lines = [str(x).strip() for x in items if x and str(x).strip()]
            if title or lines:
                other_sections.append({"title": title, "items": lines})

    deduped_other = []
    seen_other = set()
    for sec in other_sections:
        sig = (sec.get("title", "").strip().casefold(), tuple(i.casefold() for i in sec.get("items", [])))
        if sig in seen_other:
            continue
        seen_other.add(sig)
        deduped_other.append(sec)

    context["other_sections"] = deduped_other
    doc.render(context)
    try:
        doc.save(output_path)
        return output_path
    except PermissionError:
        base, ext = os.path.splitext(output_path)
        counter = 1
        while True:
            target_path = f"{base}_{counter:02d}{ext}"
            try:
                doc.save(target_path)
                return target_path
            except PermissionError:
                counter += 1





def smart_anonymize_data(data, api_key, cfg):
    blind = copy.deepcopy(data)
    in_tok, out_tok, cost = 0, 0, 0.0
    
    if cfg.get("anon_cut_name", True):
        name = blind.get('basics', {}).get('name', '')
        if name:
            parts = name.split()
            blind['basics']['name'] = f"{parts[0]} {parts[1][0]}." if len(parts) > 1 else "Candidate"
    
    if cfg.get("anon_remove_creds", True):
        if 'contacts' in blind.get('basics', {}): blind['basics']['contacts'] = {}
        if 'links' in blind.get('basics', {}): blind['basics']['links'] = []
    
    if cfg.get("anon_mask_companies", True):
        experiences = blind.get('experience', [])
        companies = [job.get('company_name') for job in experiences if job.get('company_name')]
        volunteering = blind.get('volunteering', [])
        v_orgs = [v.get('organization') for v in volunteering if v.get('organization')]
        all_companies = companies + v_orgs
        
        if all_companies:
            unique_comps = list(set(all_companies))
            
            # Use editable prompt from config
            prompt_template = cfg.get("prompt_anonymize", DEFAULT_PROMPTS["prompt_anonymize"])
            prompt = prompt_template.replace("{companies_json}", json.dumps(unique_comps, ensure_ascii=False))
            
            try:
                client = genai.Client(api_key=api_key)
                response = _retry_generate(client, MODEL_NAME, prompt)
                
                i_tok = getattr(response.usage_metadata, 'prompt_token_count', 0)
                o_tok = getattr(response.usage_metadata, 'candidates_token_count', 0)
                in_tok += i_tok
                out_tok += o_tok
                cost += (i_tok / 1_000_000 * PRICE_1M_IN) + (o_tok / 1_000_000 * PRICE_1M_OUT)
                
                text = response.text.replace('```json', '').replace('```', '').strip()
                mapping = json.loads(text)
            except: 
                mapping = {name: "Confidential Company" for name in unique_comps}

            for job in experiences:
                original = job.get('company_name')
                if original and original in mapping: job['company_name'] = mapping[original]
                elif original: job['company_name'] = "Confidential Company"
                
            for vol in volunteering:
                original = vol.get('organization')
                if original and original in mapping: vol['organization'] = mapping[original]
                elif original: vol['organization'] = "Confidential Organization"
            
    return blind, in_tok, out_tok, cost
